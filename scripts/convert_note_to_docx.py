"""
Convert Methodological Note from Markdown to Word Document
"""

import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
INPUT_FILE = os.path.join(BASE_DIR, 'reports', 'Methodological_Note.md')
OUTPUT_FILE = os.path.join(BASE_DIR, 'reports', 'Methodological_Note.docx')


def setup_styles(doc):
    """Set up document styles."""
    # Modify Normal style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Modify Heading 1
    style = doc.styles['Heading 1']
    style.font.name = 'Calibri'
    style.font.size = Pt(16)
    style.font.bold = True
    style.font.color.rgb = RGBColor(31, 78, 121)  # Dark blue

    # Modify Heading 2
    style = doc.styles['Heading 2']
    style.font.name = 'Calibri'
    style.font.size = Pt(14)
    style.font.bold = True
    style.font.color.rgb = RGBColor(46, 117, 182)  # Medium blue

    # Modify Heading 3
    style = doc.styles['Heading 3']
    style.font.name = 'Calibri'
    style.font.size = Pt(12)
    style.font.bold = True
    style.font.color.rgb = RGBColor(46, 117, 182)


def add_table_from_markdown(doc, table_lines):
    """Convert markdown table to Word table."""
    # Parse table
    rows = []
    for line in table_lines:
        if '|' in line and not line.strip().startswith('|--') and not re.match(r'^\|[-:\s|]+\|$', line.strip()):
            cells = [cell.strip() for cell in line.strip().split('|')[1:-1]]
            if cells:
                rows.append(cells)

    if not rows:
        return

    # Create table
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                cell.text = cell_text
                # Bold header row
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

    doc.add_paragraph()  # Add space after table


def add_code_block(doc, code_lines):
    """Add formatted code block."""
    para = doc.add_paragraph()
    para.style = 'Normal'
    for line in code_lines:
        run = para.add_run(line + '\n')
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
    para.paragraph_format.left_indent = Inches(0.5)


def process_inline_formatting(paragraph, text):
    """Process inline markdown formatting (bold, italic, code)."""
    # Pattern for **bold**, *italic*, `code`
    pattern = r'(\*\*.*?\*\*|\*.*?\*|`.*?`)'
    parts = re.split(pattern, text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def convert_md_to_docx():
    """Convert markdown file to Word document."""
    print(f"Reading: {INPUT_FILE}")

    with open(INPUT_FILE, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')

    doc = Document()
    setup_styles(doc)

    # Add title
    title = doc.add_heading('Methodological Note', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Health Financing Gap Statistical Product for Africa')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.italic = True

    doc.add_paragraph()  # Space

    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i]

        # Skip the main title (already added)
        if line.startswith('# Methodological Note'):
            i += 1
            continue

        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Tables
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            add_table_from_markdown(doc, table_lines)
            table_lines = []
            in_table = False

        # Headings
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=3)
            i += 1
            continue

        # Horizontal rule
        if line.strip() == '---':
            doc.add_paragraph('_' * 50)
            i += 1
            continue

        # Bullet points
        if line.strip().startswith('- '):
            para = doc.add_paragraph(style='List Bullet')
            process_inline_formatting(para, line.strip()[2:])
            i += 1
            continue

        # Numbered lists
        match = re.match(r'^(\d+)\.\s+(.+)$', line.strip())
        if match:
            para = doc.add_paragraph(style='List Number')
            process_inline_formatting(para, match.group(2))
            i += 1
            continue

        # Regular paragraphs
        if line.strip():
            para = doc.add_paragraph()
            process_inline_formatting(para, line.strip())

        i += 1

    # Handle any remaining table
    if in_table and table_lines:
        add_table_from_markdown(doc, table_lines)

    # Save document
    doc.save(OUTPUT_FILE)
    print(f"Saved: {OUTPUT_FILE}")
    print("Conversion complete!")


if __name__ == "__main__":
    convert_md_to_docx()
